#!/usr/bin/env python3
"""
Data Ingestion Service - 2025 Best Practices
Handles YouTube transcripts, books, articles, poetry, and other content types
Integrates with existing Tessera SQLite database
"""

import os
import re
import sqlite3
import asyncio
import tempfile
import mimetypes
from pathlib import Path
from typing import List, Dict, Optional, Any, Union
from datetime import datetime
from contextlib import asynccontextmanager
from urllib.parse import urlparse, parse_qs
from io import BytesIO

import structlog
import aiohttp
import aiofiles
from tenacity import retry, stop_after_attempt, wait_exponential

# Import our centralized logging
from logging_config import (
    get_logger, log_service_start, log_service_ready, 
    log_api_request, log_api_response, log_error,
    log_processing_start, log_processing_complete
)

from fastapi import FastAPI, HTTPException, Depends, BackgroundTasks, status, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field, ConfigDict, field_validator
from pydantic_settings import BaseSettings, SettingsConfigDict

# Content processing libraries
import yt_dlp
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import (
    TranscriptsDisabled, NoTranscriptFound, VideoUnavailable,
    RequestBlocked, YouTubeRequestFailed, IpBlocked
)
import PyPDF2
import docx
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
import requests
from readability import Document
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
import spacy

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')


# ========== CONFIGURATION ==========

class DataIngestionSettings(BaseSettings):
    """Modern configuration with Pydantic Settings"""
    
    model_config = SettingsConfigDict(
        env_file=".env",
        env_prefix="INGESTION_",
        case_sensitive=False,
        extra="ignore"
    )
    
    # Database
    database_path: Path = Field(
        default=Path("../data/tessera_knowledge.db"),
        description="Path to Tessera SQLite database"
    )
    
    # Service
    host: str = Field(default="127.0.0.1")
    port: int = Field(default=8003)
    
    # Processing limits
    max_file_size_mb: int = Field(default=50, description="Max file size in MB")
    max_text_length: int = Field(default=1000000, description="Max text length")
    chunk_size: int = Field(default=1000, description="Text chunk size for processing")
    
    # YouTube settings
    youtube_language_preference: List[str] = Field(default=["en", "en-US", "en-GB"])
    youtube_use_proxy: bool = Field(default=False, description="Use proxy for YouTube requests")
    youtube_proxy_url: Optional[str] = Field(default=None, description="Proxy URL for YouTube requests")
    youtube_max_retries: int = Field(default=5, description="Max retries for YouTube requests")
    youtube_retry_delay: int = Field(default=10, description="Delay between retries in seconds")
    youtube_use_cookies: bool = Field(default=False, description="Use cookies for YouTube requests")
    youtube_cookies_file: Optional[str] = Field(default=None, description="Path to cookies file")
    
    # Web scraping settings
    user_agent: str = Field(default="Tessera Data Ingestion Bot 1.0")
    request_timeout: int = Field(default=30)


# ========== PYDANTIC MODELS ==========

class ContentType(str):
    YOUTUBE = "youtube"
    BOOK = "book"
    ARTICLE = "article"
    POETRY = "poetry"
    TEXT = "text"
    PDF = "pdf"
    DOCX = "docx"
    EPUB = "epub"
    WEB = "web"


class DataIngestionRequest(BaseModel):
    model_config = ConfigDict(frozen=True)
    
    content_type: str = Field(..., description="Type of content to ingest")
    source: str = Field(..., description="URL, file path, or text content")
    title: Optional[str] = Field(None, description="Optional title override")
    description: Optional[str] = Field(None, description="Optional description")
    tags: List[str] = Field(default_factory=list, description="Content tags")
    project_id: Optional[int] = Field(None, description="Associated project ID")


class IngestionResult(BaseModel):
    model_config = ConfigDict(frozen=True)
    
    success: bool
    content_id: Optional[int] = None
    title: str
    content_type: str
    word_count: int
    chunk_count: int
    processing_time_seconds: float
    error: Optional[str] = None
    metadata: Dict[str, Any] = Field(default_factory=dict)


class HealthResponse(BaseModel):
    model_config = ConfigDict(frozen=True)
    
    status: str
    timestamp: datetime
    services: Dict[str, bool]


# ========== CONTENT PROCESSORS ==========

class YouTubeProcessor:
    """Process YouTube videos and extract transcripts"""
    
    def __init__(self, settings: DataIngestionSettings):
        self.settings = settings
        self.logger = get_logger("youtube_processor")
    
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
    async def extract_transcript(self, url: str) -> Dict[str, Any]:
        """Extract transcript from YouTube URL"""
        log_processing_start("YouTube transcript extraction", url=url)
        start_time = asyncio.get_event_loop().time()
        
        try:
            # Extract video ID from URL
            video_id = self._extract_video_id(url)
            if not video_id:
                raise ValueError(f"Could not extract video ID from URL: {url}")
            
            # Get video metadata using yt-dlp with enhanced options
            ydl_opts = {
                'quiet': True,
                'no_warnings': True,
                'extract_flat': False,
                'writesubtitles': False,
                'writeautomaticsub': False,
                'skip_download': True,
                'format': 'best[height<=480]/worst',  # Use lower quality to avoid format issues
                'user_agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                'extractor_retries': 3,
                'fragment_retries': 3,
                'retries': 3,
            }
            
            # Add proxy if configured
            if self.settings.youtube_use_proxy and self.settings.youtube_proxy_url:
                ydl_opts['proxy'] = self.settings.youtube_proxy_url
            
            # Add cookies if configured
            if self.settings.youtube_use_cookies and self.settings.youtube_cookies_file:
                ydl_opts['cookiefile'] = self.settings.youtube_cookies_file
            
            # Try to get video info with fallback options
            info = None
            for attempt in range(self.settings.youtube_max_retries):
                try:
                    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                        info = await asyncio.get_event_loop().run_in_executor(
                            None, ydl.extract_info, url, False
                        )
                    break
                except Exception as e:
                    self.logger.warning(f"yt-dlp attempt {attempt + 1} failed: {str(e)}")
                    if attempt < self.settings.youtube_max_retries - 1:
                        await asyncio.sleep(self.settings.youtube_retry_delay)
                        # Try with different format on retry
                        if attempt == 1:
                            ydl_opts['format'] = 'worst'
                        elif attempt == 2:
                            ydl_opts['format'] = 'bestaudio/best'
                    else:
                        # If all yt-dlp attempts fail, continue with basic info
                        self.logger.warning(f"All yt-dlp attempts failed, using basic video info")
                        info = {'title': f'YouTube Video {video_id}', 'description': '', 'duration': 0}
            
            # Get transcript using the correct API with enhanced error handling
            def get_transcript():
                
                # Try multiple approaches for transcript extraction
                transcript_data = None
                transcript_language = 'en'
                
                # Approach 1: Try YouTube Transcript API
                try:
                    api = YouTubeTranscriptApi()
                    transcript_list = api.list(video_id)
                    
                    # Try to get transcript in preferred language
                    transcript = None
                    for lang in self.settings.youtube_language_preference:
                        try:
                            transcript = transcript_list.find_transcript([lang])
                            break
                        except:
                            continue
                    
                    if not transcript:
                        # Get any available transcript
                        try:
                            transcript = transcript_list.find_generated_transcript(['en'])
                        except:
                            # Try any available transcript
                            available_transcripts = transcript_list._manually_created_transcripts
                            if available_transcripts:
                                transcript = list(available_transcripts.values())[0]
                            else:
                                available_transcripts = transcript_list._generated_transcripts
                                if available_transcripts:
                                    transcript = list(available_transcripts.values())[0]
                    
                    if transcript:
                        # Fetch transcript data and return both data and language
                        transcript_data = transcript.fetch()
                        transcript_language = transcript.language_code
                        return transcript_data, transcript_language
                        
                except (IpBlocked, RequestBlocked) as e:
                    # IP blocking or request blocking - try alternative approaches
                    raise e  # Re-raise to trigger retry logic
                except (TranscriptsDisabled, NoTranscriptFound, VideoUnavailable) as e:
                    # These are permanent failures
                    raise ValueError(f"Transcript not available: {str(e)}")
                except Exception as e:
                    # Other errors - log and continue to fallback
                    print(f"YouTube Transcript API failed: {str(e)}")
                
                # If we get here, transcript extraction failed
                raise ValueError("No transcript could be extracted from this video")
            
            # Try transcript extraction with retry logic for IP blocking
            transcript_data = None
            transcript_language = 'en'
            
            for attempt in range(self.settings.youtube_max_retries):
                try:
                    transcript_data, transcript_language = await asyncio.get_event_loop().run_in_executor(
                        None, get_transcript
                    )
                    break
                except Exception as e:
                    error_str = str(e).lower()
                    if 'ipblocked' in error_str or 'ip' in error_str and 'block' in error_str:
                        self.logger.warning(f"IP blocked on attempt {attempt + 1}, waiting before retry...")
                        if attempt < self.settings.youtube_max_retries - 1:
                            # Exponential backoff for IP blocking
                            wait_time = self.settings.youtube_retry_delay * (2 ** attempt)
                            await asyncio.sleep(wait_time)
                        else:
                            raise ValueError(f"YouTube transcript extraction failed after {self.settings.youtube_max_retries} attempts due to IP blocking. Try again later or use a proxy.")
                    elif 'too many requests' in error_str:
                        self.logger.warning(f"Rate limited on attempt {attempt + 1}, waiting before retry...")
                        if attempt < self.settings.youtube_max_retries - 1:
                            wait_time = self.settings.youtube_retry_delay * (3 ** attempt)  # Longer wait for rate limiting
                            await asyncio.sleep(wait_time)
                        else:
                            raise ValueError(f"YouTube transcript extraction failed due to rate limiting. Try again later.")
                    else:
                        # Other errors - shorter retry
                        self.logger.warning(f"Transcript extraction attempt {attempt + 1} failed: {str(e)}")
                        if attempt < self.settings.youtube_max_retries - 1:
                            await asyncio.sleep(5)
                        else:
                            raise ValueError(f"Failed to extract transcript: {str(e)}")
            
            # Combine transcript text
            full_text = " ".join([entry.text for entry in transcript_data])
            
            duration_ms = (asyncio.get_event_loop().time() - start_time) * 1000
            log_processing_complete("YouTube transcript extraction", duration_ms, 
                                  video_id=video_id, transcript_length=len(full_text))
            
            # Ensure we have some content
            if not full_text or len(full_text.strip()) < 10:
                raise ValueError("Extracted transcript is too short or empty")
            
            return {
                'title': info.get('title', f'YouTube Video {video_id}'),
                'description': info.get('description', ''),
                'duration': info.get('duration', 0),
                'uploader': info.get('uploader', ''),
                'upload_date': info.get('upload_date', ''),
                'view_count': info.get('view_count', 0),
                'transcript': full_text,
                'transcript_language': transcript_language,
                'url': url,
                'video_id': video_id,
                'extraction_method': 'youtube_transcript_api'
            }
            
        except ValueError as e:
            # These are expected errors with user-friendly messages
            log_error(e, "YouTube transcript extraction failed", url=url)
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=str(e)
            )
        except Exception as e:
            # Unexpected errors
            error_msg = str(e)
            if 'ipblocked' in error_msg.lower() or ('ip' in error_msg.lower() and 'block' in error_msg.lower()):
                user_msg = "YouTube has temporarily blocked requests from your IP address. This usually happens due to too many requests. Please try again later or consider using a proxy."
            elif 'too many requests' in error_msg.lower() or 'rate limit' in error_msg.lower():
                user_msg = "YouTube is rate limiting requests. Please wait a few minutes before trying again."
            elif 'forbidden' in error_msg.lower() or '403' in error_msg:
                user_msg = "Access to this YouTube video is restricted. The video may be private, age-restricted, or geo-blocked."
            elif 'not available' in error_msg.lower() or 'unavailable' in error_msg.lower():
                user_msg = "This YouTube video is not available or has been removed."
            else:
                user_msg = f"Failed to extract YouTube transcript: {error_msg}"
            
            log_error(e, "Failed to extract YouTube transcript", url=url)
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=user_msg
            )
    
    def _extract_video_id(self, url: str) -> Optional[str]:
        """Extract YouTube video ID from various URL formats"""
        patterns = [
            r'(?:youtube\.com/watch\?v=|youtu\.be/|youtube\.com/embed/)([^&\n?#]+)',
            r'youtube\.com/v/([^&\n?#]+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)
        
        return None


class BookProcessor:
    """Process books and documents (PDF, DOCX, EPUB, TXT)"""
    
    def __init__(self, settings: DataIngestionSettings):
        self.settings = settings
        self.logger = get_logger("book_processor")
    
    async def process_file(self, file_content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """Process uploaded file and extract text"""
        try:
            # Determine file type
            file_extension = Path(filename).suffix.lower()
            
            if file_extension == '.pdf':
                text = await self._extract_pdf_text(file_content)
            elif file_extension == '.docx':
                text = await self._extract_docx_text(file_content)
            elif file_extension == '.epub':
                text = await self._extract_epub_text(file_content)
            elif file_extension in ['.txt', '.md']:
                text = file_content.decode('utf-8', errors='ignore')
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Clean and process text
            text = self._clean_text(text)
            
            return {
                'title': Path(filename).stem,
                'content': text,
                'file_type': file_extension,
                'file_size': len(file_content),
                'word_count': len(text.split()),
                'filename': filename
            }
            
        except Exception as e:
            self.logger.error("Failed to process book file", filename=filename, error=str(e))
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to process file: {str(e)}"
            )
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF"""
        def extract():
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        
        return await asyncio.get_event_loop().run_in_executor(None, extract)
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX"""
        def extract():
            doc = docx.Document(BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        return await asyncio.get_event_loop().run_in_executor(None, extract)
    
    async def _extract_epub_text(self, file_content: bytes) -> str:
        """Extract text from EPUB"""
        def extract():
            book = epub.read_epub(BytesIO(file_content))
            text = ""
            
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    text += soup.get_text() + "\n"
            
            return text
        
        return await asyncio.get_event_loop().run_in_executor(None, extract)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
        return text.strip()


class ArticleProcessor:
    """Process web articles and news content"""
    
    def __init__(self, settings: DataIngestionSettings):
        self.settings = settings
        self.logger = get_logger("article_processor")
    
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
    async def extract_article(self, url: str) -> Dict[str, Any]:
        """Extract article content from URL"""
        try:
            headers = {
                'User-Agent': self.settings.user_agent,
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.5',
                'Accept-Encoding': 'gzip, deflate',
                'Connection': 'keep-alive',
            }
            
            async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=self.settings.request_timeout)) as session:
                async with session.get(url, headers=headers) as response:
                    if response.status != 200:
                        raise HTTPException(
                            status_code=status.HTTP_400_BAD_REQUEST,
                            detail=f"Failed to fetch article: HTTP {response.status}"
                        )
                    
                    html_content = await response.text()
            
            # Use readability to extract main content
            doc = Document(html_content)
            
            # Parse with BeautifulSoup for additional metadata
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract metadata
            title = doc.title() or self._extract_title_from_soup(soup)
            content = doc.summary()
            
            # Clean content
            content_soup = BeautifulSoup(content, 'html.parser')
            clean_content = content_soup.get_text()
            clean_content = self._clean_text(clean_content)
            
            # Extract additional metadata
            description = self._extract_meta_description(soup)
            author = self._extract_author(soup)
            publish_date = self._extract_publish_date(soup)
            
            return {
                'title': title,
                'content': clean_content,
                'description': description,
                'author': author,
                'publish_date': publish_date,
                'url': url,
                'word_count': len(clean_content.split()),
                'domain': urlparse(url).netloc
            }
            
        except Exception as e:
            self.logger.error("Failed to extract article", url=url, error=str(e))
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to extract article: {str(e)}"
            )
    
    def _extract_title_from_soup(self, soup: BeautifulSoup) -> str:
        """Extract title from HTML soup"""
        title_tag = soup.find('title')
        if title_tag:
            return title_tag.get_text().strip()
        
        h1_tag = soup.find('h1')
        if h1_tag:
            return h1_tag.get_text().strip()
        
        return "Unknown Article"
    
    def _extract_meta_description(self, soup: BeautifulSoup) -> str:
        """Extract meta description"""
        meta_desc = soup.find('meta', attrs={'name': 'description'})
        if meta_desc:
            return meta_desc.get('content', '').strip()
        
        meta_desc = soup.find('meta', attrs={'property': 'og:description'})
        if meta_desc:
            return meta_desc.get('content', '').strip()
        
        return ""
    
    def _extract_author(self, soup: BeautifulSoup) -> str:
        """Extract author information"""
        # Try various author meta tags
        author_selectors = [
            'meta[name="author"]',
            'meta[property="article:author"]',
            'meta[name="article:author"]',
            '.author',
            '.byline',
            '[rel="author"]'
        ]
        
        for selector in author_selectors:
            element = soup.select_one(selector)
            if element:
                if element.name == 'meta':
                    return element.get('content', '').strip()
                else:
                    return element.get_text().strip()
        
        return ""
    
    def _extract_publish_date(self, soup: BeautifulSoup) -> str:
        """Extract publish date"""
        date_selectors = [
            'meta[property="article:published_time"]',
            'meta[name="article:published_time"]',
            'meta[name="date"]',
            'time[datetime]',
            '.date',
            '.published'
        ]
        
        for selector in date_selectors:
            element = soup.select_one(selector)
            if element:
                if element.name == 'meta':
                    return element.get('content', '').strip()
                elif element.name == 'time':
                    return element.get('datetime', element.get_text()).strip()
                else:
                    return element.get_text().strip()
        
        return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
        return text.strip()


class PoetryProcessor:
    """Process poetry and creative writing"""
    
    def __init__(self, settings: DataIngestionSettings):
        self.settings = settings
        self.logger = get_logger("poetry_processor")
        
        # Load spaCy model for NLP processing
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            self.logger.warning("spaCy model 'en_core_web_sm' not found. Poetry analysis will be limited.")
            self.nlp = None
    
    async def process_poetry(self, text: str, title: Optional[str] = None) -> Dict[str, Any]:
        """Process poetry text and extract literary features"""
        try:
            # Clean text
            clean_text = self._clean_text(text)
            
            # Analyze structure
            lines = [line.strip() for line in clean_text.split('\n') if line.strip()]
            stanzas = self._identify_stanzas(lines)
            
            # Basic metrics
            word_count = len(clean_text.split())
            line_count = len(lines)
            stanza_count = len(stanzas)
            
            # Literary analysis
            analysis = {}
            if self.nlp:
                analysis = await self._analyze_literary_features(clean_text)
            
            return {
                'title': title or "Untitled Poem",
                'content': clean_text,
                'lines': lines,
                'stanzas': stanzas,
                'word_count': word_count,
                'line_count': line_count,
                'stanza_count': stanza_count,
                'analysis': analysis,
                'content_type': 'poetry'
            }
            
        except Exception as e:
            self.logger.error("Failed to process poetry", error=str(e))
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to process poetry: {str(e)}"
            )
    
    def _identify_stanzas(self, lines: List[str]) -> List[List[str]]:
        """Identify stanzas in poetry"""
        stanzas = []
        current_stanza = []
        
        for line in lines:
            if line.strip():
                current_stanza.append(line)
            else:
                if current_stanza:
                    stanzas.append(current_stanza)
                    current_stanza = []
        
        if current_stanza:
            stanzas.append(current_stanza)
        
        return stanzas
    
    async def _analyze_literary_features(self, text: str) -> Dict[str, Any]:
        """Analyze literary features using NLP"""
        def analyze():
            doc = self.nlp(text)
            
            # Extract entities
            entities = [(ent.text, ent.label_) for ent in doc.ents]
            
            # Extract key phrases (noun phrases)
            noun_phrases = [chunk.text for chunk in doc.noun_chunks]
            
            # Sentiment analysis (basic)
            sentiment_score = sum([token.sentiment for token in doc if hasattr(token, 'sentiment')]) / len(doc)
            
            # POS tags
            pos_counts = {}
            for token in doc:
                pos = token.pos_
                pos_counts[pos] = pos_counts.get(pos, 0) + 1
            
            return {
                'entities': entities,
                'noun_phrases': noun_phrases[:10],  # Top 10
                'sentiment_score': sentiment_score,
                'pos_distribution': pos_counts
            }
        
        return await asyncio.get_event_loop().run_in_executor(None, analyze)
    
    def _clean_text(self, text: str) -> str:
        """Clean text while preserving poetry structure"""
        # Remove control characters but preserve line breaks
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
        return text.strip()


# ========== DATABASE MANAGER ==========

class DatabaseManager:
    """Manage database operations for ingested content"""
    
    def __init__(self, settings: DataIngestionSettings):
        self.settings = settings
        self.logger = structlog.get_logger(__name__)
    
    async def store_content(self, content_data: Dict[str, Any], content_type: str, project_id: Optional[int] = None) -> int:
        """Store processed content in database"""
        try:
            db_path = str(self.settings.database_path.resolve())
            
            def store():
                conn = sqlite3.connect(db_path)
                cursor = conn.cursor()
                
                # Insert into articles table (matching actual schema)
                cursor.execute("""
                    INSERT INTO articles (
                        title, url, content, summary, categories, 
                        images, sections, infobox, coordinates
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    content_data.get('title', 'Untitled'),
                    content_data.get('url', ''),
                    content_data.get('content', ''),
                    content_data.get('description', ''),
                    '[]',  # categories as JSON
                    '[]',  # images as JSON
                    '[]',  # sections as JSON
                    '{}',  # infobox as JSON
                    None   # coordinates
                ))
                
                article_id = cursor.lastrowid
                
                # Also insert into learning_content table for dashboard integration
                cursor.execute("""
                    INSERT INTO learning_content (
                        title, content_type, url, content, summary, 
                        metadata, source, difficulty_level, estimated_time_minutes,
                        completion_percentage, tags
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    content_data.get('title', 'Untitled'),
                    content_type,
                    content_data.get('url', ''),
                    content_data.get('content', ''),
                    content_data.get('description', ''),
                    '{}',  # metadata as JSON
                    'ingested',  # source
                    self._estimate_difficulty(content_data),
                    self._estimate_time(content_data),
                    0,  # completion_percentage starts at 0
                    '[]'  # tags as JSON array
                ))
                
                learning_content_id = cursor.lastrowid
                
                # Link to appropriate subject based on content type and analysis
                subject_id = self._determine_subject(content_data, content_type)
                if subject_id:
                    cursor.execute("""
                        INSERT INTO content_subjects (content_id, subject_id, relevance_score)
                        VALUES (?, ?, ?)
                    """, (learning_content_id, subject_id, 1.0))
                
                conn.commit()
                conn.close()
                
                return article_id
            
            return await asyncio.get_event_loop().run_in_executor(None, store)
            
        except Exception as e:
            self.logger.error("Failed to store content", error=str(e))
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to store content: {str(e)}"
            )
    
    def _estimate_difficulty(self, content_data: Dict[str, Any]) -> int:
        """Estimate content difficulty on 1-5 scale"""
        content = content_data.get('content', '')
        word_count = len(content.split()) if content else 0
        
        # Simple heuristic based on content length and complexity
        if word_count < 500:
            return 1  # Easy
        elif word_count < 2000:
            return 2  # Medium-Easy
        elif word_count < 5000:
            return 3  # Medium
        elif word_count < 10000:
            return 4  # Hard
        else:
            return 5  # Very Hard
    
    def _estimate_time(self, content_data: Dict[str, Any]) -> int:
        """Estimate reading/consumption time in minutes"""
        content = content_data.get('content', '')
        word_count = len(content.split()) if content else 0
        
        # Average reading speed: 200 words per minute
        reading_time = max(1, word_count // 200)
        
        # Add extra time for videos (if duration is available)
        if 'duration' in content_data and content_data['duration']:
            return content_data['duration'] // 60  # Convert seconds to minutes
        
        return reading_time
    
    def _determine_subject(self, content_data: Dict[str, Any], content_type: str) -> Optional[int]:
        """Determine which subject this content belongs to based on content analysis"""
        title = content_data.get('title', '').lower()
        content = content_data.get('content', '').lower()
        
        # Simple keyword-based subject mapping
        # In a real system, you'd use NLP/ML for better classification
        
        programming_keywords = ['python', 'javascript', 'code', 'programming', 'software', 'algorithm', 'function', 'variable']
        ml_keywords = ['machine learning', 'ai', 'artificial intelligence', 'neural network', 'deep learning', 'data science']
        web_keywords = ['html', 'css', 'react', 'vue', 'angular', 'web development', 'frontend', 'backend']
        cooking_keywords = ['recipe', 'cooking', 'food', 'kitchen', 'ingredient', 'chef', 'culinary']
        data_science_keywords = ['data analysis', 'statistics', 'visualization', 'pandas', 'numpy', 'analytics']
        personal_dev_keywords = ['productivity', 'self-improvement', 'motivation', 'habits', 'goals', 'personal development']
        
        text_to_analyze = f"{title} {content}"
        
        # Count keyword matches for each subject
        subjects = {
            1: programming_keywords,    # Programming
            2: cooking_keywords,        # Cooking  
            3: ml_keywords,            # Machine Learning
            4: web_keywords,           # Web Development
            5: data_science_keywords,  # Data Science
            6: personal_dev_keywords   # Personal Development
        }
        
        best_subject = None
        max_matches = 0
        
        for subject_id, keywords in subjects.items():
            matches = sum(1 for keyword in keywords if keyword in text_to_analyze)
            if matches > max_matches:
                max_matches = matches
                best_subject = subject_id
        
        # Default to Programming if no clear match and it's code-related content
        if not best_subject and content_type in ['youtube', 'article']:
            if any(keyword in text_to_analyze for keyword in ['tech', 'tutorial', 'guide', 'how to']):
                best_subject = 1  # Programming as default for tech content
        
        return best_subject


# ========== MAIN SERVICE ==========

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Application lifespan manager"""
    log_service_start("Data Ingestion Service", port=8001)
    
    # Initialize services
    settings = DataIngestionSettings()
    
    # Store in app state
    app.state.settings = settings
    app.state.youtube_processor = YouTubeProcessor(settings)
    app.state.book_processor = BookProcessor(settings)
    app.state.article_processor = ArticleProcessor(settings)
    app.state.poetry_processor = PoetryProcessor(settings)
    app.state.db_manager = DatabaseManager(settings)
    
    yield
    
    logger = get_logger("system")
    logger.info("🛑 Shutting down Data Ingestion Service")


# Create FastAPI app
app = FastAPI(
    title="Tessera Data Ingestion Service",
    description="Advanced content ingestion for YouTube, books, articles, and poetry",
    version="1.0.0",
    lifespan=lifespan
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure structured logging
structlog.configure(
    processors=[
        structlog.stdlib.filter_by_level,
        structlog.stdlib.add_logger_name,
        structlog.stdlib.add_log_level,
        structlog.stdlib.PositionalArgumentsFormatter(),
        structlog.processors.TimeStamper(fmt="iso"),
        structlog.processors.StackInfoRenderer(),
        structlog.processors.format_exc_info,
        structlog.processors.UnicodeDecoder(),
        structlog.processors.JSONRenderer()
    ],
    context_class=dict,
    logger_factory=structlog.stdlib.LoggerFactory(),
    wrapper_class=structlog.stdlib.BoundLogger,
    cache_logger_on_first_use=True,
)


# ========== API ENDPOINTS ==========

@app.get("/health", response_model=HealthResponse)
async def health_check():
    """Health check endpoint"""
    return HealthResponse(
        status="healthy",
        timestamp=datetime.now(),
        services={
            "database": True,  # Could add actual DB check
            "youtube": True,
            "article_extraction": True,
            "text_processing": True
        }
    )


@app.post("/ingest/youtube", response_model=IngestionResult)
async def ingest_youtube(
    url: str = Form(...),
    title: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    project_id: Optional[int] = Form(None),
    background_tasks: BackgroundTasks = BackgroundTasks()
):
    log_api_request("POST", "/ingest/youtube", url=url)
    """Ingest YouTube video transcript"""
    start_time = asyncio.get_event_loop().time()
    
    try:
        youtube_data = await app.state.youtube_processor.extract_transcript(url)
        
        # Store in database
        content_id = await app.state.db_manager.store_content(
            youtube_data, ContentType.YOUTUBE, project_id
        )
        
        # Create chunks for embedding (background task)
        background_tasks.add_task(
            create_embeddings_for_content, content_id, youtube_data['transcript']
        )
        
        processing_time = asyncio.get_event_loop().time() - start_time
        duration_ms = processing_time * 1000
        log_api_response("POST", "/ingest/youtube", 200, duration_ms, 
                        content_id=content_id, title=youtube_data['title'])
        
        return IngestionResult(
            success=True,
            content_id=content_id,
            title=youtube_data['title'],
            content_type=ContentType.YOUTUBE,
            word_count=len(youtube_data['transcript'].split()),
            chunk_count=len(youtube_data['transcript']) // app.state.settings.chunk_size + 1,
            processing_time_seconds=processing_time,
            metadata={
                'duration': youtube_data['duration'],
                'uploader': youtube_data['uploader'],
                'view_count': youtube_data['view_count']
            }
        )
        
    except Exception as e:
        processing_time = asyncio.get_event_loop().time() - start_time
        return IngestionResult(
            success=False,
            title=title or "Failed YouTube Video",
            content_type=ContentType.YOUTUBE,
            word_count=0,
            chunk_count=0,
            processing_time_seconds=processing_time,
            error=str(e)
        )


@app.post("/ingest/book", response_model=IngestionResult)
async def ingest_book(
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    project_id: Optional[int] = Form(None),
    background_tasks: BackgroundTasks = BackgroundTasks()
):
    """Ingest book or document file"""
    start_time = asyncio.get_event_loop().time()
    
    try:
        # Check file size
        file_content = await file.read()
        if len(file_content) > app.state.settings.max_file_size_mb * 1024 * 1024:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"File too large. Max size: {app.state.settings.max_file_size_mb}MB"
            )
        
        # Process file
        book_data = await app.state.book_processor.process_file(
            file_content, file.filename, file.content_type
        )
        
        # Override title if provided
        if title:
            book_data['title'] = title
        if description:
            book_data['description'] = description
        
        # Store in database
        content_id = await app.state.db_manager.store_content(
            book_data, ContentType.BOOK, project_id
        )
        
        # Create chunks for embedding (background task)
        background_tasks.add_task(
            create_embeddings_for_content, content_id, book_data['content']
        )
        
        processing_time = asyncio.get_event_loop().time() - start_time
        
        return IngestionResult(
            success=True,
            content_id=content_id,
            title=book_data['title'],
            content_type=ContentType.BOOK,
            word_count=book_data['word_count'],
            chunk_count=len(book_data['content']) // app.state.settings.chunk_size + 1,
            processing_time_seconds=processing_time,
            metadata={
                'file_type': book_data['file_type'],
                'file_size': book_data['file_size'],
                'filename': book_data['filename']
            }
        )
        
    except Exception as e:
        processing_time = asyncio.get_event_loop().time() - start_time
        return IngestionResult(
            success=False,
            title=title or file.filename or "Failed Book Upload",
            content_type=ContentType.BOOK,
            word_count=0,
            chunk_count=0,
            processing_time_seconds=processing_time,
            error=str(e)
        )


@app.post("/ingest/article", response_model=IngestionResult)
async def ingest_article(
    url: str = Form(...),
    title: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    project_id: Optional[int] = Form(None),
    background_tasks: BackgroundTasks = BackgroundTasks()
):
    """Ingest web article"""
    start_time = asyncio.get_event_loop().time()
    
    try:
        article_data = await app.state.article_processor.extract_article(url)
        
        # Override title/description if provided
        if title:
            article_data['title'] = title
        if description:
            article_data['description'] = description
        
        # Store in database
        content_id = await app.state.db_manager.store_content(
            article_data, ContentType.ARTICLE, project_id
        )
        
        # Create chunks for embedding (background task)
        background_tasks.add_task(
            create_embeddings_for_content, content_id, article_data['content']
        )
        
        processing_time = asyncio.get_event_loop().time() - start_time
        
        return IngestionResult(
            success=True,
            content_id=content_id,
            title=article_data['title'],
            content_type=ContentType.ARTICLE,
            word_count=article_data['word_count'],
            chunk_count=len(article_data['content']) // app.state.settings.chunk_size + 1,
            processing_time_seconds=processing_time,
            metadata={
                'author': article_data['author'],
                'domain': article_data['domain'],
                'publish_date': article_data['publish_date']
            }
        )
        
    except Exception as e:
        processing_time = asyncio.get_event_loop().time() - start_time
        return IngestionResult(
            success=False,
            title=title or "Failed Article",
            content_type=ContentType.ARTICLE,
            word_count=0,
            chunk_count=0,
            processing_time_seconds=processing_time,
            error=str(e)
        )


@app.post("/ingest/poetry", response_model=IngestionResult)
async def ingest_poetry(
    text: str = Form(...),
    title: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    project_id: Optional[int] = Form(None),
    background_tasks: BackgroundTasks = BackgroundTasks()
):
    """Ingest poetry or creative writing"""
    start_time = asyncio.get_event_loop().time()
    
    try:
        poetry_data = await app.state.poetry_processor.process_poetry(text, title)
        
        if description:
            poetry_data['description'] = description
        
        # Store in database
        content_id = await app.state.db_manager.store_content(
            poetry_data, ContentType.POETRY, project_id
        )
        
        # Create chunks for embedding (background task)
        background_tasks.add_task(
            create_embeddings_for_content, content_id, poetry_data['content']
        )
        
        processing_time = asyncio.get_event_loop().time() - start_time
        
        return IngestionResult(
            success=True,
            content_id=content_id,
            title=poetry_data['title'],
            content_type=ContentType.POETRY,
            word_count=poetry_data['word_count'],
            chunk_count=poetry_data['stanza_count'],  # Use stanzas as chunks for poetry
            processing_time_seconds=processing_time,
            metadata={
                'line_count': poetry_data['line_count'],
                'stanza_count': poetry_data['stanza_count'],
                'analysis': poetry_data.get('analysis', {})
            }
        )
        
    except Exception as e:
        processing_time = asyncio.get_event_loop().time() - start_time
        return IngestionResult(
            success=False,
            title=title or "Failed Poetry",
            content_type=ContentType.POETRY,
            word_count=0,
            chunk_count=0,
            processing_time_seconds=processing_time,
            error=str(e)
        )


# ========== BACKGROUND TASKS ==========

async def create_embeddings_for_content(content_id: int, text: str):
    """Background task to create embeddings for ingested content"""
    try:
        # This would call the embedding service
        # For now, just log that we would create embeddings
        logger = structlog.get_logger(__name__)
        logger.info("Creating embeddings for content", content_id=content_id, text_length=len(text))
        
        # TODO: Call embedding service API
        # await call_embedding_service(content_id, text)
        
    except Exception as e:
        logger = structlog.get_logger(__name__)
        logger.error("Failed to create embeddings", content_id=content_id, error=str(e))


# ========== MAIN ==========

if __name__ == "__main__":
    import uvicorn
    
    settings = DataIngestionSettings()
    
    uvicorn.run(
        "data_ingestion_service:app",
        host=settings.host,
        port=settings.port,
        reload=True,
        log_level="info"
    )
